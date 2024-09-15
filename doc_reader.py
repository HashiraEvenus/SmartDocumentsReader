from docx import Document
import PyPDF2
import os
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter

def read_file(file_path):
    _, file_extension = os.path.splitext(file_path)
    
    if file_extension.lower() == '.docx':
        return read_docx(file_path)
    elif file_extension.lower() == '.pdf':
        return read_pdf(file_path)
    elif file_extension.lower() == '.txt':
        return read_txt(file_path)
    else:
        raise ValueError(f"Unsupported file format: {file_extension}")

def read_docx(file_path):
    doc = Document(file_path)
    return '\n'.join([paragraph.text for paragraph in doc.paragraphs])

def read_pdf(file_path):
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        return '\n'.join([page.extract_text() for page in reader.pages])

def read_txt(file_path):
    with open(file_path, 'r', encoding='utf-8') as file:
        return file.read()

def save_file(file_path, content):
    _, file_extension = os.path.splitext(file_path)
    
    if file_extension.lower() == '.docx':
        save_docx(file_path, content)
    elif file_extension.lower() in ['.txt', '.pdf']:
        save_txt(file_path, content)
    else:
        raise ValueError(f"Unsupported file format for saving: {file_extension}")

def save_docx(file_path, content):
    doc = Document()
    for paragraph in content.split('\n'):
        doc.add_paragraph(paragraph)
    doc.save(file_path)

def save_txt(file_path, content):
    with open(file_path, 'w', encoding='utf-8') as file:
        file.write(content)

def export_to_pdf(file_path, content):
    c = canvas.Canvas(file_path, pagesize=letter)
    width, height = letter
    y = height - 50  # Start from top of the page
    for line in content.split('\n'):
        if y < 50:  # If we're at the bottom of the page
            c.showPage()  # Start a new page
            y = height - 50  # Reset y to top of page
        c.drawString(50, y, line)
        y -= 15  # Move down for next line
    c.save()

def export_to_html(file_path, content):
    html_content = f"""
    <!DOCTYPE html>
    <html>
    <head>
        <title>Exported Document</title>
    </head>
    <body>
        <pre>{content}</pre>
    </body>
    </html>
    """
    with open(file_path, 'w', encoding='utf-8') as f:
        f.write(html_content)

def extract_toc(file_path):
    _, file_extension = os.path.splitext(file_path)
    
    if file_extension.lower() == '.docx':
        return extract_toc_docx(file_path)
    elif file_extension.lower() == '.pdf':
        return extract_toc_pdf(file_path)
    else:
        return []  # Return empty list for unsupported formats

def extract_toc_docx(file_path):
    doc = Document(file_path)
    toc = []
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith('Heading'):
            level = int(paragraph.style.name[-1])
            toc.append({'title': paragraph.text, 'level': level, 'page': 0})  # Page number not available in docx
    return toc

def extract_toc_pdf(file_path):
    with open(file_path, 'rb') as file:
        reader = PyPDF2.PdfReader(file)
        if '/Outlines' in reader.trailer['/Root']:
            outlines = reader.outline
            return parse_pdf_outlines(outlines)
    return []

def parse_pdf_outlines(outlines, level=1):
    toc = []
    for item in outlines:
        if isinstance(item, list):
            toc.extend(parse_pdf_outlines(item, level+1))
        else:
            toc.append({'title': item.title, 'level': level, 'page': item.page})
    return toc

def generate_bullet_points(text, num_points=5):
    summary = summarize_text(text, num_points)
    sentences = sent_tokenize(summary)
    return '\n'.join([f"• {sentence}" for sentence in sentences])

def summarize_text(text, num_sentences=5):
    # Tokenize the text into sentences
    sentences = sent_tokenize(text)
    
    # Remove stopwords and punctuation
    stop_words = set(stopwords.words('english') + list('!"#$%&\'()*+,-./:;<=>?@[\\]^_`{|}~'))
    words = [word.lower() for sentence in sentences for word in sentence.split() if word.lower() not in stop_words]
    
    # Calculate word frequencies
    freq = FreqDist(words)
    
    # Score sentences based on word frequencies
    sentence_scores = {}
    for sentence in sentences:
        for word in sentence.split():
            if word.lower() in freq:
                if sentence not in sentence_scores:
                    sentence_scores[sentence] = freq[word.lower()]
                else:
                    sentence_scores[sentence] += freq[word.lower()]
    
    # Get the top N sentences
    summary_sentences = nlargest(num_sentences, sentence_scores, key=sentence_scores.get)
    
    # Join the sentences and return the summary
    return ' '.join(summary_sentences)





